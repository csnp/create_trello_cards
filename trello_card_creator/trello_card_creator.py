#!/usr/bin/env python3


import os
import sys
import re
import datetime
import getpass
import logging
import argparse
from typing import List, Optional, Dict, Any

import requests
import docx
import keyring
from tkinter import Tk
from tkinter.filedialog import askopenfilename

# Configure logging
logging.basicConfig(level=logging.INFO, format="%(levelname)s: %(message)s")

# Constants for keyring
KEYRING_SERVICE = "Trello_Card_Creator"
API_KEY_ENTRY = "api_key"
API_TOKEN_ENTRY = "api_token"
BOARD_ID_ENTRY = "board_id"

# Mapping from list names to list IDs (caches list IDs to avoid multiple API calls)
list_name_to_id: Dict[str, str] = {}


class TrelloCardCreatorError(Exception):
    """Custom exception for Trello Card Creator errors."""

    pass


def load_credentials_and_board_id() -> (str, str, str):
    """Load Trello API credentials and Board ID securely or prompt the user to enter them."""
    # Check if credentials are stored
    api_key = keyring.get_password(KEYRING_SERVICE, API_KEY_ENTRY)
    api_token = keyring.get_password(KEYRING_SERVICE, API_TOKEN_ENTRY)
    board_id = keyring.get_password(KEYRING_SERVICE, BOARD_ID_ENTRY)

    if api_key and api_token and board_id:
        update = (
            input("Do you want to update your API Key, Token, or Board ID? (y/n): ")
            .strip()
            .lower()
        )
        if update == "y":
            # Prompt for new values
            api_key, api_token, board_id = prompt_for_new_credentials(
                api_key, api_token, board_id
            )
        else:
            logging.info("Using stored API credentials and Board ID.")
            # Verify access to the board
            if not verify_board_access(board_id, api_key, api_token):
                logging.error(
                    "Access to the stored Board ID is denied. Please update your API Token and Board ID."
                )
                api_key, api_token, board_id = prompt_for_new_credentials(
                    api_key, api_token, board_id, force_token_update=True
                )
    else:
        api_key, api_token, board_id = prompt_for_new_credentials(
            api_key, api_token, board_id
        )
    return api_key, api_token, board_id


def prompt_for_new_credentials(
    stored_api_key: Optional[str],
    stored_api_token: Optional[str],
    stored_board_id: Optional[str],
    force_token_update: bool = False,
) -> (str, str, str):
    """Prompt the user to enter new API credentials and Board ID."""
    print("\nPlease enter new values. Leave blank to keep the existing value.\n")

    # Prompt for API Key
    api_key_prompt = "Enter your Trello API Key [{}]: ".format(
        "********" if stored_api_key else ""
    )
    api_key = input(api_key_prompt).strip()
    if not api_key:
        api_key = stored_api_key

    # Prompt for API Token
    if force_token_update:
        print("You need to enter a new API Token.")
        api_token = getpass.getpass(
            "Enter your Trello API Token (input is hidden): "
        ).strip()
    else:
        api_token_prompt = "Enter your Trello API Token [{}]: ".format(
            "********" if stored_api_token else ""
        )
        api_token = getpass.getpass(api_token_prompt).strip()
        if not api_token:
            api_token = stored_api_token

    # Determine if API Token was updated
    token_updated = api_token != stored_api_token

    # Prompt for Board ID
    if token_updated or not stored_board_id:
        print(
            "\nSince the API Token was updated, you need to provide the Trello Board URL."
        )
        board_id = prompt_for_board_id(api_key, api_token)
    else:
        board_url_prompt = "Enter your Trello Board URL [{}]: ".format("stored")
        board_url = input(board_url_prompt).strip()
        if not board_url:
            board_id = stored_board_id
        else:
            board_identifier = extract_board_id_from_url(board_url)
            if not board_identifier:
                raise TrelloCardCreatorError(
                    "Unable to extract or find the board ID from the URL."
                )
            board_id = get_board_id(board_identifier, api_key, api_token)

    # Store updated values securely
    keyring.set_password(KEYRING_SERVICE, API_KEY_ENTRY, api_key)
    keyring.set_password(KEYRING_SERVICE, API_TOKEN_ENTRY, api_token)
    keyring.set_password(KEYRING_SERVICE, BOARD_ID_ENTRY, board_id)
    logging.info("Credentials updated and saved securely.\n")
    return api_key, api_token, board_id


def prompt_for_board_id(api_key: str, api_token: str) -> str:
    """Prompt the user to enter the Trello board URL and return the Board ID."""
    board_url = input(
        "Enter your Trello board URL (e.g., https://trello.com/b/YourBoardID): "
    ).strip()
    board_identifier = extract_board_id_from_url(board_url)
    if not board_identifier:
        raise TrelloCardCreatorError(
            "Unable to extract or find the board ID from the URL."
        )
    # Get the full board ID
    board_id = get_board_id(board_identifier, api_key, api_token)
    return board_id


def extract_board_id_from_url(board_url: str) -> Optional[str]:
    """Extract the board ID or short link from the Trello board URL."""
    pattern = r"trello\.com/(b|board)/([\w\d]+)/?"
    match = re.search(pattern, board_url)
    if match:
        board_short_id = match.group(2)
        return board_short_id
    else:
        logging.error(
            "Could not extract board ID from URL. Please ensure the URL is correct."
        )
        return None


def get_board_id(board_identifier: str, api_key: str, api_token: str) -> str:
    """Get the full board ID using the board identifier (ID or short link)."""
    url = f"https://api.trello.com/1/boards/{board_identifier}"
    query = {"key": api_key, "token": api_token, "fields": "id"}
    response = requests.get(url, params=query)
    if response.status_code == 200:
        board = response.json()
        # Verify access to the board
        if verify_board_access(board["id"], api_key, api_token):
            return board["id"]
        else:
            error_message = (
                "Access to the specified board is denied. Please ensure your API token has access to this board.\n"
                "Refer to the documentation on how to generate a token with appropriate access:\n"
                "https://github.com/csnp/create_trello_cards\n"
            )
            raise TrelloCardCreatorError(error_message)
    else:
        raise TrelloCardCreatorError(f"Error fetching board ID: {response.text}")


def verify_board_access(board_id: str, api_key: str, api_token: str) -> bool:
    """Verify if the API credentials have access to the specified board."""
    url = f"https://api.trello.com/1/boards/{board_id}"
    query = {"key": api_key, "token": api_token, "fields": "id"}
    response = requests.get(url, params=query)
    if response.status_code == 200:
        return True
    elif response.status_code == 401:
        # Unauthorized access
        return False
    else:
        logging.error(f"Error verifying board access: {response.text}")
        return False


def get_list_id(
    list_name: str, board_id: str, api_key: str, api_token: str
) -> Optional[str]:
    """Get the ID of a Trello list by its name. Create it if it doesn't exist."""
    if list_name in list_name_to_id:
        return list_name_to_id[list_name]
    else:
        # Fetch list IDs from Trello
        url = f"https://api.trello.com/1/boards/{board_id}/lists"
        query = {"key": api_key, "token": api_token, "fields": "name,id"}
        response = requests.get(url, params=query)
        if response.status_code == 200:
            lists = response.json()
            for trello_list in lists:
                list_name_to_id[trello_list["name"]] = trello_list["id"]
            if list_name in list_name_to_id:
                return list_name_to_id[list_name]
            else:
                # List not found, create it
                logging.info(
                    f"List '{list_name}' not found on the Trello board. Creating it now."
                )
                new_list_id = create_list(list_name, board_id, api_key, api_token)
                if new_list_id:
                    list_name_to_id[list_name] = new_list_id
                    return new_list_id
                else:
                    logging.error(f"Failed to create list '{list_name}'.")
                    return None
        else:
            logging.error(f"Error fetching lists from Trello: {response.text}")
            return None


def create_list(
    list_name: str, board_id: str, api_key: str, api_token: str
) -> Optional[str]:
    """Create a new list on the Trello board."""
    url = f"https://api.trello.com/1/lists"
    query = {"key": api_key, "token": api_token, "name": list_name, "idBoard": board_id}
    response = requests.post(url, params=query)
    if response.status_code == 200:
        trello_list = response.json()
        logging.info(f"List '{list_name}' created successfully.")
        return trello_list["id"]
    else:
        logging.error(f"Error creating list '{list_name}': {response.text}")
        return None


def get_label_id(
    label_name: str, board_id: str, api_key: str, api_token: str
) -> Optional[str]:
    """Get the ID of a Trello label by its name, or create it if it doesn't exist."""
    # Fetch labels from Trello board
    url = f"https://api.trello.com/1/boards/{board_id}/labels"
    query = {"key": api_key, "token": api_token, "fields": "name,id"}
    response = requests.get(url, params=query)
    if response.status_code == 200:
        labels = response.json()
        for label in labels:
            if label["name"].lower() == label_name.lower():
                return label["id"]
        # If label not found, create it
        create_label_url = f"https://api.trello.com/1/labels"
        query = {
            "key": api_key,
            "token": api_token,
            "idBoard": board_id,
            "name": label_name,
            "color": None,  # You can assign a color if desired
        }
        response = requests.post(create_label_url, params=query)
        if response.status_code == 200:
            label = response.json()
            return label["id"]
        else:
            logging.error(f"Error creating label '{label_name}': {response.text}")
            return None
    else:
        logging.error(f"Error fetching labels from Trello: {response.text}")
        return None


def get_member_id(
    member_identifier: str, api_key: str, api_token: str
) -> Optional[str]:
    """Get the ID of a Trello member by username."""
    url = f"https://api.trello.com/1/members/{member_identifier}"
    query = {"key": api_key, "token": api_token}
    response = requests.get(url, params=query)
    if response.status_code == 200:
        member = response.json()
        return member["id"]
    else:
        logging.error(f"Error fetching member '{member_identifier}': {response.text}")
        return None


def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    doc = docx.Document(file_path)
    cards = []
    paragraphs = doc.paragraphs

    i = 0
    while i < len(paragraphs):
        para = paragraphs[i]
        if para.style.name == "Heading 1":
            # Start of a new card
            card = {
                "title": para.text.strip(),
                "description": "",
                "labels": [],
                "due_date": None,
                "members": [],
                "list_name": None,
                "checklist": [],
                "attachments": [],
                "image": None,
            }
            i += 1

            # Read description until we find a special field or a new Heading 1
            while i < len(paragraphs) and paragraphs[i].style.name != "Heading 1":
                text = paragraphs[i].text.strip()
                if text.startswith(
                    (
                        "Labels:",
                        "Due Date:",
                        "Members:",
                        "List:",
                        "Checklist:",
                        "Attachments:",
                        "Image:",
                    )
                ):
                    break
                else:
                    card["description"] += text + "\n"
                    i += 1
            card["description"] = card["description"].strip()

            # Now parse optional fields
            while i < len(paragraphs) and paragraphs[i].style.name != "Heading 1":
                text = paragraphs[i].text.strip()
                if text.startswith("Labels:"):
                    labels = text[7:].strip().split(",")
                    card["labels"] = [
                        label.strip() for label in labels if label.strip()
                    ]
                    i += 1
                elif text.startswith("Due Date:"):
                    card["due_date"] = text[9:].strip()
                    i += 1
                elif text.startswith("Members:"):
                    members = text[8:].strip().split(",")
                    card["members"] = [
                        member.strip() for member in members if member.strip()
                    ]
                    i += 1
                elif text.startswith("List:"):
                    card["list_name"] = text[5:].strip()
                    i += 1
                elif text.startswith("Checklist:"):
                    i += 1
                    while i < len(paragraphs):
                        p = paragraphs[i]
                        if p.style.name in ["List Bullet", "List Number"]:
                            card["checklist"].append(p.text.strip())
                            i += 1
                        elif p.style.name == "Normal" and p.text.strip() == "":
                            i += 1
                        else:
                            break
                elif text.startswith("Attachments:"):
                    i += 1
                    while i < len(paragraphs):
                        p = paragraphs[i]
                        if p.text.strip() != "" and not p.text.strip().startswith(
                            "Image:"
                        ):
                            card["attachments"].append(p.text.strip())
                            i += 1
                        else:
                            break
                elif text.startswith("Image:"):
                    card["image"] = text[6:].strip()
                    i += 1
                else:
                    i += 1

            cards.append(card)
        else:
            i += 1

    return cards


def create_trello_card(
    card: Dict[str, Any], board_id: str, api_key: str, api_token: str
) -> bool:
    """Create a Trello card with the specified properties."""
    # Determine the list ID
    if "list_name" in card and card["list_name"]:
        list_id = get_list_id(card["list_name"], board_id, api_key, api_token)
        if not list_id:
            logging.error(f"Skipping card '{card['title']}' due to missing list.")
            return False
    else:
        # Default to the first list on the board
        url = f"https://api.trello.com/1/boards/{board_id}/lists"
        query = {"key": api_key, "token": api_token, "limit": 1}
        response = requests.get(url, params=query)
        if response.status_code == 200:
            lists = response.json()
            if lists:
                list_id = lists[0]["id"]
            else:
                logging.error(
                    f"No lists found on the board. Cannot create card '{card['title']}'."
                )
                return False
        else:
            logging.error(f"Error fetching lists: {response.text}")
            return False
    # Prepare the card data
    card_data = {
        "name": card["title"],
        "desc": card["description"],
        "idList": list_id,
        "key": api_key,
        "token": api_token,
    }
    # Handle due date
    if "due_date" in card and card["due_date"]:
        try:
            due_date = card["due_date"]
            # Trello expects ISO 8601 format with time and timezone
            due_date_iso = datetime.datetime.fromisoformat(due_date).isoformat()
            card_data["due"] = due_date_iso
        except ValueError:
            logging.warning(
                f"Invalid due date format for card '{card['title']}': {card['due_date']}"
            )
    # Handle labels
    label_ids = []
    if "labels" in card and card["labels"]:
        for label_name in card["labels"]:
            label_id = get_label_id(label_name, board_id, api_key, api_token)
            if label_id:
                label_ids.append(label_id)
        if label_ids:
            card_data["idLabels"] = ",".join(label_ids)
    # Handle members
    member_ids = []
    if "members" in card and card["members"]:
        for member_identifier in card["members"]:
            member_id = get_member_id(member_identifier, api_key, api_token)
            if member_id:
                member_ids.append(member_id)
        if member_ids:
            card_data["idMembers"] = ",".join(member_ids)
    # Create the card
    url = "https://api.trello.com/1/cards"
    response = requests.post(url, params=card_data)
    if response.status_code == 200:
        created_card = response.json()
        logging.info(f"Card '{card['title']}' created successfully.")
        # Handle checklist
        if "checklist" in card and card["checklist"]:
            create_checklist(created_card["id"], card["checklist"], api_key, api_token)
        # Handle attachments
        if "attachments" in card and card["attachments"]:
            for attachment in card["attachments"]:
                add_attachment(created_card["id"], attachment, api_key, api_token)
        # Handle image (setting as cover image)
        if "image" in card and card["image"]:
            add_attachment(
                created_card["id"], card["image"], api_key, api_token, set_cover=True
            )
        return True
    else:
        logging.error(f"Error creating card '{card['title']}': {response.text}")
        return False


def create_checklist(card_id: str, items: List[str], api_key: str, api_token: str):
    """Create a checklist on a Trello card."""
    url = f"https://api.trello.com/1/cards/{card_id}/checklists"
    query = {"key": api_key, "token": api_token, "name": "Checklist"}
    response = requests.post(url, params=query)
    if response.status_code == 200:
        checklist = response.json()
        checklist_id = checklist["id"]
        for item in items:
            add_checklist_item(checklist_id, item, api_key, api_token)
    else:
        logging.error(f"Error creating checklist: {response.text}")


def add_checklist_item(checklist_id: str, item_name: str, api_key: str, api_token: str):
    """Add an item to a checklist."""
    url = f"https://api.trello.com/1/checklists/{checklist_id}/checkItems"
    query = {"key": api_key, "token": api_token, "name": item_name}
    response = requests.post(url, params=query)
    if response.status_code != 200:
        logging.error(f"Error adding checklist item '{item_name}': {response.text}")


def add_attachment(
    card_id: str,
    attachment_url_or_path: str,
    api_key: str,
    api_token: str,
    set_cover: bool = False,
):
    """Add an attachment to a Trello card."""
    url = f"https://api.trello.com/1/cards/{card_id}/attachments"
    if re.match(r"^https?://", attachment_url_or_path):
        # It's a URL
        query = {
            "key": api_key,
            "token": api_token,
            "url": attachment_url_or_path,
            "setCover": str(set_cover).lower(),
        }
        response = requests.post(url, params=query)
    else:
        # It's a file path
        if os.path.isfile(attachment_url_or_path):
            with open(attachment_url_or_path, "rb") as file_data:
                files = {"file": file_data}
                query = {
                    "key": api_key,
                    "token": api_token,
                    "setCover": str(set_cover).lower(),
                }
                response = requests.post(url, params=query, files=files)
        else:
            logging.warning(f"Attachment file not found: {attachment_url_or_path}")
            return
    if response.status_code != 200:
        logging.error(
            f"Error adding attachment '{attachment_url_or_path}': {response.text}"
        )


def generate_sample_docx():
    """Generate a sample DOCX template for users to fill in."""
    from docx import Document

    doc = Document()
    doc.add_heading("Sample Card Title", level=1)
    doc.add_paragraph("This is the description of the card.")
    doc.add_paragraph("Labels: Marketing, Urgent")
    doc.add_paragraph("Due Date: 2023-12-31T23:59:00")
    doc.add_paragraph("Members: username1, username2")
    doc.add_paragraph("List: To Do")
    doc.add_paragraph("Checklist:", style="Normal")
    doc.add_paragraph("First item", style="List Bullet")
    doc.add_paragraph("Second item", style="List Bullet")
    doc.add_paragraph("Attachments:", style="Normal")
    doc.add_paragraph("https://example.com/document.pdf")
    doc.add_paragraph("/path/to/file.pdf")
    doc.add_paragraph("Image: https://example.com/image.png")
    sample_file = "sample_template.docx"
    doc.save(sample_file)
    logging.info(
        f"\nSample template '{sample_file}' has been created in the current directory."
    )
    logging.info(
        "Please fill in your card details in this file and run the script again."
    )
    sys.exit(0)


def main():
    """Main function to execute the script."""
    print("Welcome to the Trello Card Creator Script!")
    print("-----------------------------------------\n")

    # Load or prompt for API credentials and Board ID
    api_key, api_token, board_id = load_credentials_and_board_id()

    # Parse command-line arguments
    parser = argparse.ArgumentParser(
        description="Create Trello cards from a DOCX file."
    )
    parser.add_argument("--file", help="Path to the DOCX file.")
    args = parser.parse_args()

    if args.file:
        file_path = args.file
        if not os.path.isfile(file_path):
            logging.error(f"The file '{file_path}' does not exist.")
            sys.exit(1)
    else:
        # Prompt the user to select a DOCX file
        print("\nPlease select your DOCX file containing the card details.")
        print(
            "If you don't have one, the script can generate a sample template for you."
        )
        choice = input("Do you have a DOCX file ready? (y/n): ").strip().lower()
        if choice == "y":
            # Use Tkinter to open a file dialog
            Tk().withdraw()  # We don't want a full GUI, so keep the root window from appearing
            file_path = askopenfilename(
                title="Select DOCX File", filetypes=[("Word Documents", "*.docx")]
            )
            if not file_path:
                logging.error("No file selected. Exiting.")
                sys.exit(1)
        else:
            generate_sample_docx()

    # Parse the DOCX file and create Trello cards
    cards = parse_docx(file_path)
    if not cards:
        logging.error(
            "\nNo cards found in the DOCX file. Please ensure the file is formatted correctly."
        )
        logging.error(
            "Each card should start with a Heading 1 style for the card title."
        )
        sys.exit(1)

    print("\nCreating Trello cards...")
    total_cards = len(cards)
    created_count = 0
    failed_count = 0

    for card in cards:
        success = create_trello_card(card, board_id, api_key, api_token)
        if success:
            created_count += 1
        else:
            failed_count += 1

    logging.info(
        f"\nProcessing complete. {created_count} out of {total_cards} cards were created successfully."
    )
    if failed_count > 0:
        logging.warning(
            f"{failed_count} cards could not be created. Please review the error messages above."
        )


if __name__ == "__main__":
    try:
        main()
    except TrelloCardCreatorError as e:
        logging.error(e)
        sys.exit(1)
    except KeyboardInterrupt:
        logging.info("\nScript interrupted by user.")
        sys.exit(0)
