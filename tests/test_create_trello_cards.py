import sys
import os
import unittest
from unittest.mock import patch, MagicMock
import requests
import docx

# Add the parent directory to the Python path
sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from trello_card_creator.trello_card_creator import (
    extract_board_id_from_url,
    get_board_id,
    verify_board_access,
    TrelloCardCreatorError,
    parse_docx,
)


class TestCreateTrelloCards(unittest.TestCase):

    def test_extract_board_id_from_valid_url(self):
        url = "https://trello.com/b/abcd1234/board-name"
        expected = "abcd1234"
        result = extract_board_id_from_url(url)
        self.assertEqual(result, expected)

    @patch("trello_card_creator.trello_card_creator.logging.error")
    def test_extract_board_id_from_invalid_url(self, mock_error):
        url = "https://trello.com/b//board-name"
        result = extract_board_id_from_url(url)
        self.assertIsNone(result)
        mock_error.assert_called_once_with(
            "Could not extract board ID from URL. Please ensure the URL is correct."
        )

    def test_extract_board_id_from_non_trello_url(self):
        url = "https://example.com/b/abcd1234/board-name"
        result = extract_board_id_from_url(url)
        self.assertIsNone(result)

    @patch("requests.get")
    def test_get_board_id_success(self, mock_get):
        mock_response = MagicMock()
        mock_response.status_code = 200
        mock_response.json.return_value = {"id": "mock_board_id"}
        mock_get.return_value = mock_response

        with patch(
            "trello_card_creator.trello_card_creator.verify_board_access",
            return_value=True,
        ):
            board_id = get_board_id("abcd1234", "fake_api_key", "fake_api_token")
            self.assertEqual(board_id, "mock_board_id")

    @patch("requests.get")
    def test_get_board_id_failure(self, mock_get):
        mock_response = MagicMock()
        mock_response.status_code = 404
        mock_response.text = "Board not found"
        mock_get.return_value = mock_response

        with self.assertRaises(TrelloCardCreatorError):
            get_board_id("invalid_id", "fake_api_key", "fake_api_token")

    @patch("requests.get")
    def test_verify_board_access_success(self, mock_get):
        mock_response = MagicMock()
        mock_response.status_code = 200
        mock_get.return_value = mock_response

        access = verify_board_access("mock_board_id", "fake_api_key", "fake_api_token")
        self.assertTrue(access)

    @patch("requests.get")
    def test_verify_board_access_failure(self, mock_get):
        mock_response = MagicMock()
        mock_response.status_code = 401
        mock_get.return_value = mock_response

        access = verify_board_access("mock_board_id", "fake_api_key", "fake_api_token")
        self.assertFalse(access)

    def test_parse_docx(self):
        # Create a mock document
        from docx import Document
        from docx.enum.style import WD_STYLE_TYPE
        from io import BytesIO

        doc = Document()
        doc.add_heading("Test Card Title", level=1)
        doc.add_paragraph("This is a test description.")
        doc.add_paragraph("Labels: TestLabel1, TestLabel2")
        doc.add_paragraph("Due Date: 2023-12-31T23:59:00")
        doc.add_paragraph("Members: testuser1, testuser2")
        doc.add_paragraph("List: Test List")
        doc.add_paragraph("Checklist:", style="Normal")
        doc.add_paragraph("Item 1", style="List Bullet")
        doc.add_paragraph("Item 2", style="List Bullet")
        doc.add_paragraph("Attachments:", style="Normal")
        doc.add_paragraph("https://example.com/file.pdf")
        doc.add_paragraph("Image: https://example.com/image.png")

        # Save the document to a bytes buffer
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Read the document from the bytes buffer
        test_doc = docx.Document(doc_io)
        # Mock the docx.Document call to return our test_doc
        with patch("docx.Document", return_value=test_doc):
            cards = parse_docx("dummy_path.docx")
            self.assertEqual(len(cards), 1)
            card = cards[0]
            self.assertEqual(card["title"], "Test Card Title")
            self.assertEqual(card["description"], "This is a test description.")
            self.assertEqual(card["labels"], ["TestLabel1", "TestLabel2"])
            self.assertEqual(card["due_date"], "2023-12-31T23:59:00")
            self.assertEqual(card["members"], ["testuser1", "testuser2"])
            self.assertEqual(card["list_name"], "Test List")
            self.assertEqual(card["checklist"], ["Item 1", "Item 2"])
            self.assertEqual(card["attachments"], ["https://example.com/file.pdf"])
            self.assertEqual(card["image"], "https://example.com/image.png")


if __name__ == "__main__":
    unittest.main()
